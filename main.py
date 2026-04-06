"""EvalBot CLI — run the multi-agent startup evaluation pipeline."""

from __future__ import annotations

from enum import Enum
import json
import os
import re
import shutil
import sys
import threading
import time
from datetime import datetime
from pathlib import Path
from typing import Any, get_args, get_origin

import pdfplumber
from dotenv import load_dotenv

load_dotenv()

from src.docs import load_submission
from src.db import get_retry_stats

PROJECT_ROOT = Path(__file__).resolve().parent

# ANSI color codes for terminal output
class Colors:
    RESET = "\033[0m"
    BOLD = "\033[1m"
    DIM = "\033[2m"
    
    # Regular colors
    BLACK = "\033[30m"
    RED = "\033[31m"
    GREEN = "\033[32m"
    YELLOW = "\033[33m"
    BLUE = "\033[34m"
    MAGENTA = "\033[35m"
    CYAN = "\033[36m"
    WHITE = "\033[37m"
    
    # Bright colors
    BRIGHT_RED = "\033[91m"
    BRIGHT_GREEN = "\033[92m"
    BRIGHT_YELLOW = "\033[93m"
    BRIGHT_BLUE = "\033[94m"
    BRIGHT_MAGENTA = "\033[95m"
    BRIGHT_CYAN = "\033[96m"
    
    # Background colors
    BG_BLUE = "\033[44m"
    BG_MAGENTA = "\033[45m"
    BG_CYAN = "\033[46m"


def _colorize(text: str, color: str) -> str:
    """Apply ANSI color to text."""
    return f"{color}{text}{Colors.RESET}"


def _next_batch_id() -> str:
    """Return the next sequential batch ID (batch_1, batch_2, ...)."""
    output_dir = PROJECT_ROOT / "output" / "Batch"
    if not output_dir.exists():
        return "batch_1"
    existing = [
        int(d.name.split("_")[1])
        for d in output_dir.iterdir()
        if d.is_dir() and d.name.startswith("batch_") and d.name.split("_")[1].isdigit()
    ]
    return f"batch_{max(existing) + 1}" if existing else "batch_1"


def _ensure_supported_python() -> None:
    """CrewAI stack used by this project is not compatible with Python 3.14+."""
    if sys.version_info < (3, 14):
        return

    py313 = PROJECT_ROOT / ".venv313" / "bin" / "python"

    if py313.exists() and Path(sys.executable).resolve() != py313.resolve():
        print("Python 3.14 detected; re-launching with .venv313 for compatibility...\n")
        os.execv(str(py313), [str(py313), str(Path(__file__).resolve()), *sys.argv[1:]])

    print("EvalBot currently requires Python 3.13 or earlier.")
    print("Create .venv313 and run with: .venv313/bin/python main.py")
    sys.exit(1)


def _extract_startup_name(text: str) -> str:
    """Try to extract a startup name from submission text, fall back to generic."""
    for line in text.splitlines():
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped[:80]
    return "Unknown Startup"


def _sanitize_filename(name: str) -> str:
    """Make a string safe for use as a filename."""
    return "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()


def _get_pdf_reference(pdf_path: Path) -> str:
    """Return a file reference marker for PDF that Agent 1 (Claude) can process directly."""
    return f"[PDF_FILE: {pdf_path.absolute()}]"


def _extract_text_from_docx(docx_path: Path) -> str | None:
    """Extract text and tables from DOCX file, preserving document order.
    
    Extracts:
    - All paragraph text
    - All tables (formatted as pipe-separated rows)
    - Content in document order (paragraphs and tables intermixed)
    
    Does NOT extract:
    - Images (stored as separate binary files)
    - Headers/footers (in separate XML files)
    - Complex formatting (bold, italics, colors)
    
    Returns None if extraction fails.
    """
    try:
        import docx
        doc = docx.Document(docx_path)
        parts = []
        
        # Process all body elements in document order
        # This preserves the natural flow of paragraphs and tables
        for element in doc.element.body:
            # Paragraph element
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                        break
            
            # Table element
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        # Format table as pipe-separated rows
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells)
                            if row_text.strip():  # Skip empty rows
                                parts.append(row_text)
                        parts.append("")  # Blank line after table
                        break
        
        return "\n".join(parts) if parts else None
    except Exception:
        return None


AGENT_ROLES = {
    1: "Intake Parser",
    2: "Venture Analyst",
    3: "Market & Competition Analyst",
    4: "Product & Positioning Analyst",
    5: "Founder Fit Analyst",
    6: "Recommendation / Pivot Agent",
    7: "Ranking Committee Agent",
}


def _supports_structured_output(model_name: str) -> bool:
    """Return whether we should use output_pydantic for this model/provider."""
    return not model_name.lower().startswith("minimax/")


def _extract_json_segment(text: str) -> str | None:
    """Extract the first balanced JSON object or array from text."""
    starts = [idx for idx in (text.find("{"), text.find("[")) if idx != -1]
    if not starts:
        return None
    start = min(starts)

    stack: list[str] = []
    in_string = False
    escaped = False

    for i in range(start, len(text)):
        ch = text[i]
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
            continue
        if ch in "{[":
            stack.append(ch)
            continue
        if ch in "}]":
            if not stack:
                return None
            opener = stack.pop()
            if (opener == "{" and ch != "}") or (opener == "[" and ch != "]"):
                return None
            if not stack:
                return text[start : i + 1]

    return None


def _parse_output_to_dict(output_model: Any, raw: str) -> dict[str, Any] | None:
    """Parse LLM raw output into the expected dict, handling fenced JSON and extra prose."""
    def _normalize_for_model(data: dict[str, Any]) -> dict[str, Any]:
        def _unwrap_optional(annotation: Any) -> Any:
            origin = get_origin(annotation)
            if origin is None:
                return annotation
            args = [a for a in get_args(annotation) if a is not type(None)]
            return args[0] if len(args) == 1 else annotation

        def _coerce_rerun_agent(value: Any) -> Any:
            if value is None:
                return None
            if isinstance(value, bool):
                return None if value is False else 1
            if isinstance(value, int):
                return value
            if isinstance(value, str):
                s = value.strip().lower()
                if s in {"", "none", "null", "n/a", "na", "false", "no"}:
                    return None
                if s.isdigit():
                    return int(s)
                # Handles variants like "Agent1", "agent_2", "Agent 3 Parser".
                match = re.search(r"agent\s*[_-]?(\d+)", s)
                if match:
                    return int(match.group(1))
            return value

        def _coerce_str(value: Any, fallback: str = "Not provided") -> str:
            if value is None:
                return fallback
            if isinstance(value, str):
                s = value.strip()
                return s if s else fallback
            return str(value)

        def _clamp_score(field_name: str, value: int) -> int:
            if field_name.startswith("score_"):
                return max(1, min(10, value))
            if field_name in {"clarity_score", "fit_score", "execution_score", "attractiveness_score", "competition_score"}:
                return max(1, min(10, value))
            return value

        def _default_for_field(field_name: str, ann: Any) -> Any:
            origin = get_origin(ann)
            if origin is list:
                return []
            if origin is dict:
                return {}
            if isinstance(ann, type) and issubclass(ann, Enum):
                return next(iter(ann)).value
            if isinstance(ann, type) and ann is int:
                return 5 if "score" in field_name else 0
            if isinstance(ann, type) and ann is float:
                return 0.0
            if isinstance(ann, type) and ann is str:
                return "Not provided"
            if isinstance(ann, type) and hasattr(ann, "model_fields"):
                return _salvage_for_model(ann, {})
            return None

        def _salvage_for_model(model_cls: Any, src: dict[str, Any]) -> dict[str, Any]:
            out: dict[str, Any] = {}
            for f_name, f_info in model_cls.model_fields.items():
                f_ann = _unwrap_optional(f_info.annotation)
                f_origin = get_origin(f_ann)
                has_value = f_name in src and src[f_name] is not None
                raw_value = src.get(f_name) if has_value else None

                if not has_value:
                    if f_info.is_required():
                        out[f_name] = _default_for_field(f_name, f_ann)
                    continue

                if f_name == "rerun_from_agent":
                    out[f_name] = _coerce_rerun_agent(raw_value)
                    continue

                if isinstance(f_ann, type) and issubclass(f_ann, Enum):
                    if isinstance(raw_value, str):
                        raw_val = raw_value.strip()
                        compact = re.sub(r"[^a-z0-9]+", "", raw_val.lower())
                        matched = None
                        for enum_member in f_ann:
                            enum_val = str(enum_member.value)
                            enum_compact = re.sub(r"[^a-z0-9]+", "", enum_val.lower())
                            if raw_val == enum_val or compact == enum_compact:
                                matched = enum_member.value
                                break
                        out[f_name] = matched if matched is not None else next(iter(f_ann)).value
                    else:
                        out[f_name] = next(iter(f_ann)).value
                    continue

                if f_ann is int:
                    out[f_name] = _clamp_score(f_name, _coerce_int(raw_value) if _coerce_int(raw_value) is not None else 0)
                    continue

                if f_ann is float:
                    try:
                        out[f_name] = float(raw_value)
                    except Exception:
                        out[f_name] = 0.0
                    continue

                if f_ann is str:
                    out[f_name] = _coerce_str(raw_value)
                    continue

                if f_origin is list:
                    if isinstance(raw_value, list):
                        out[f_name] = [str(x).strip() for x in raw_value if str(x).strip()]
                    elif isinstance(raw_value, str):
                        lines = [
                            line.strip().lstrip("-*•0123456789. ").strip()
                            for line in raw_value.replace("\r", "").split("\n")
                        ]
                        parts = [p for p in lines if p]
                        if len(parts) <= 1 and "," in raw_value:
                            parts = [p.strip() for p in raw_value.split(",") if p.strip()]
                        out[f_name] = parts
                    else:
                        out[f_name] = [str(raw_value)] if raw_value is not None else []
                    continue

                if isinstance(f_ann, type) and hasattr(f_ann, "model_fields"):
                    if isinstance(raw_value, str):
                        seg = _extract_json_segment(raw_value)
                        candidate = (seg or raw_value).strip()
                        try:
                            parsed_obj = json.loads(candidate)
                        except Exception:
                            parsed_obj = {}
                        if isinstance(parsed_obj, dict):
                            out[f_name] = _salvage_for_model(f_ann, parsed_obj)
                        else:
                            out[f_name] = _salvage_for_model(f_ann, {})
                    elif isinstance(raw_value, dict):
                        out[f_name] = _salvage_for_model(f_ann, raw_value)
                    else:
                        out[f_name] = _salvage_for_model(f_ann, {})
                    continue

                out[f_name] = raw_value

            return out

        def _coerce_int(value: Any) -> Any:
            if isinstance(value, bool):
                return int(value)
            if isinstance(value, int):
                return value
            if isinstance(value, float):
                return int(round(value))
            if isinstance(value, str):
                s = value.strip()
                if not s:
                    return value
                # Handles "7", "7.0", "7/10", "Score: 7", etc.
                match = re.search(r"-?\d+(?:\.\d+)?", s)
                if match:
                    try:
                        num = float(match.group(0))
                        return int(round(num))
                    except Exception:
                        return value
            return value

        normalized = dict(data)
        for field_name, field_info in output_model.model_fields.items():
            if field_name not in normalized:
                continue
            value = normalized[field_name]
            if value is None:
                continue

            ann = _unwrap_optional(field_info.annotation)
            ann_origin = get_origin(ann)
            is_list_field = ann_origin is list

            if field_name == "rerun_from_agent":
                normalized[field_name] = _coerce_rerun_agent(value)
                continue

            # Normalize enum values with forgiving case/format matching.
            if isinstance(ann, type) and issubclass(ann, Enum) and isinstance(value, str):
                raw_val = value.strip()
                compact = re.sub(r"[^a-z0-9]+", "", raw_val.lower())
                matched = None
                for enum_member in ann:
                    enum_val = str(enum_member.value)
                    enum_compact = re.sub(r"[^a-z0-9]+", "", enum_val.lower())
                    if raw_val == enum_val or compact == enum_compact:
                        matched = enum_member.value
                        break
                if matched is not None:
                    normalized[field_name] = matched
                    continue

            # Coerce integer-like fields from strings/floats.
            if ann is int:
                normalized[field_name] = _coerce_int(value)
                continue

            # Coerce string fields from non-string values.
            if ann is str and not isinstance(value, str):
                if isinstance(value, (int, float, bool)):
                    normalized[field_name] = str(value)
                    continue
                if isinstance(value, list):
                    normalized[field_name] = "; ".join(str(v).strip() for v in value if str(v).strip())
                    continue
                if isinstance(value, dict):
                    normalized[field_name] = json.dumps(value, default=str)
                    continue

            # If a dict/model field arrives as a JSON string, parse it.
            if ann_origin in {dict} and isinstance(value, str):
                seg = _extract_json_segment(value)
                candidate = (seg or value).strip()
                try:
                    parsed_obj = json.loads(candidate)
                    if isinstance(parsed_obj, dict):
                        normalized[field_name] = parsed_obj
                        continue
                except Exception:
                    pass

            # If a nested model field arrives as a JSON string, parse it to dict first.
            if isinstance(ann, type) and hasattr(ann, "model_fields") and isinstance(value, str):
                seg = _extract_json_segment(value)
                candidate = (seg or value).strip()
                try:
                    parsed_obj = json.loads(candidate)
                    if isinstance(parsed_obj, dict):
                        normalized[field_name] = parsed_obj
                        continue
                except Exception:
                    pass

            if is_list_field:
                if isinstance(value, str):
                    lines = [
                        line.strip().lstrip("-*•0123456789. ").strip()
                        for line in value.replace("\r", "").split("\n")
                    ]
                    parts = [p for p in lines if p]
                    if len(parts) <= 1 and "," in value:
                        parts = [p.strip() for p in value.split(",") if p.strip()]
                    normalized[field_name] = parts if parts else [value.strip()]
                elif isinstance(value, tuple | set):
                    normalized[field_name] = list(value)

        return normalized

    def _salvage_for_model(model_cls: Any, src: dict[str, Any]) -> dict[str, Any]:
        def _unwrap_optional(annotation: Any) -> Any:
            origin = get_origin(annotation)
            if origin is None:
                return annotation
            args = [a for a in get_args(annotation) if a is not type(None)]
            return args[0] if len(args) == 1 else annotation

        def _fallback_for(field_name: str, ann: Any) -> Any:
            origin = get_origin(ann)
            if origin is list:
                return []
            if origin is dict:
                return {}
            if isinstance(ann, type) and issubclass(ann, Enum):
                return next(iter(ann)).value
            if ann is int:
                return 5 if "score" in field_name else 0
            if ann is float:
                return 0.0
            if ann is str:
                return "Not provided"
            if isinstance(ann, type) and hasattr(ann, "model_fields"):
                return _salvage_for_model(ann, {})
            return None

        out: dict[str, Any] = {}
        for field_name, field_info in model_cls.model_fields.items():
            ann = _unwrap_optional(field_info.annotation)
            value = src.get(field_name)

            if value is None:
                if field_info.is_required():
                    out[field_name] = _fallback_for(field_name, ann)
                continue

            if ann is int:
                if isinstance(value, int):
                    out[field_name] = value
                elif isinstance(value, float):
                    out[field_name] = int(round(value))
                elif isinstance(value, str):
                    match = re.search(r"-?\d+(?:\.\d+)?", value)
                    out[field_name] = int(round(float(match.group(0)))) if match else _fallback_for(field_name, ann)
                else:
                    out[field_name] = _fallback_for(field_name, ann)
                if field_name.startswith("score_") or field_name in {
                    "clarity_score", "fit_score", "execution_score", "attractiveness_score", "competition_score"
                }:
                    out[field_name] = max(1, min(10, int(out[field_name])))
                continue

            if ann is float:
                try:
                    out[field_name] = float(value)
                except Exception:
                    out[field_name] = 0.0
                continue

            if ann is str:
                out[field_name] = value.strip() if isinstance(value, str) and value.strip() else str(value)
                continue

            origin = get_origin(ann)
            if origin is list:
                if isinstance(value, list):
                    out[field_name] = [str(v).strip() for v in value if str(v).strip()]
                elif isinstance(value, str):
                    parts = [p.strip() for p in re.split(r"\n|,", value) if p.strip()]
                    out[field_name] = parts
                else:
                    out[field_name] = []
                continue

            if isinstance(ann, type) and issubclass(ann, Enum):
                if isinstance(value, str):
                    compact = re.sub(r"[^a-z0-9]+", "", value.lower())
                    matched = None
                    for enum_member in ann:
                        enum_compact = re.sub(r"[^a-z0-9]+", "", str(enum_member.value).lower())
                        if compact == enum_compact or value == enum_member.value:
                            matched = enum_member.value
                            break
                    out[field_name] = matched if matched is not None else next(iter(ann)).value
                else:
                    out[field_name] = next(iter(ann)).value
                continue

            if isinstance(ann, type) and hasattr(ann, "model_fields"):
                if isinstance(value, dict):
                    out[field_name] = _salvage_for_model(ann, value)
                elif isinstance(value, str):
                    seg = _extract_json_segment(value)
                    if seg:
                        try:
                            parsed_obj = json.loads(seg)
                            out[field_name] = _salvage_for_model(ann, parsed_obj if isinstance(parsed_obj, dict) else {})
                        except Exception:
                            out[field_name] = _salvage_for_model(ann, {})
                    else:
                        out[field_name] = _salvage_for_model(ann, {})
                else:
                    out[field_name] = _salvage_for_model(ann, {})
                continue

            out[field_name] = value

        return out

    candidates: list[str] = [raw.strip()]

    for block in re.findall(r"```(?:json)?\s*([\s\S]*?)\s*```", raw, flags=re.IGNORECASE):
        stripped = block.strip()
        if stripped:
            candidates.append(stripped)

    segment = _extract_json_segment(raw)
    if segment:
        candidates.append(segment.strip())

    seen: set[str] = set()
    unique_candidates: list[str] = []
    for candidate in candidates:
        if candidate and candidate not in seen:
            seen.add(candidate)
            unique_candidates.append(candidate)

    for candidate in unique_candidates:
        try:
            parsed = output_model.model_validate_json(candidate)
            return parsed.model_dump(mode="json")
        except Exception:
            pass

        try:
            loaded = json.loads(candidate)
            if isinstance(loaded, list) and loaded and isinstance(loaded[0], dict):
                loaded = loaded[0]
            if isinstance(loaded, dict):
                normalized = _normalize_for_model(loaded)
                try:
                    parsed = output_model.model_validate(normalized)
                    return parsed.model_dump(mode="json")
                except Exception:
                    try:
                        salvaged = _salvage_for_model(output_model, normalized)
                        parsed = output_model.model_validate(salvaged)
                        return parsed.model_dump(mode="json")
                    except Exception:
                        continue
        except Exception:
            continue

    return None


def _parse_failure_reason(output_model: Any, raw: str) -> str | None:
    """Return a short reason why raw output failed schema validation."""
    candidates: list[str] = [raw.strip()]

    for block in re.findall(r"```(?:json)?\s*([\s\S]*?)\s*```", raw, flags=re.IGNORECASE):
        stripped = block.strip()
        if stripped:
            candidates.append(stripped)

    segment = _extract_json_segment(raw)
    if segment:
        candidates.append(segment.strip())

    seen: set[str] = set()
    for candidate in candidates:
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)

        # First try JSON string validation directly.
        try:
            output_model.model_validate_json(candidate)
            return None
        except Exception as exc:
            msg = str(exc)

        # Then try loading and validating as dict/list.
        try:
            loaded = json.loads(candidate)
            if isinstance(loaded, list) and loaded and isinstance(loaded[0], dict):
                loaded = loaded[0]
            if isinstance(loaded, dict):
                try:
                    output_model.model_validate(loaded)
                    return None
                except Exception as exc2:
                    msg = str(exc2)
            else:
                msg = "Top-level JSON is not an object."
        except Exception as exc3:
            msg = f"JSON decode error: {exc3}"

        msg = " ".join(msg.split())
        if len(msg) > 220:
            msg = msg[:217] + "..."
        return msg

    return "No JSON object candidate found in model output."


def _repair_output_to_json(agent: Any, output_model: Any, raw: str) -> dict[str, Any] | None:
    """Ask the model to convert its own raw output into strict JSON once parsing fails."""
    from crewai import Crew as _Crew, Task as _Task

    field_names = ", ".join(output_model.model_fields.keys())
    repair_prompt = (
        "Convert the following content into exactly one valid JSON object that matches the target schema.\n"
        "Output ONLY JSON. No markdown, no code fences, no commentary.\n"
        f"Required top-level keys: {field_names}.\n\n"
        "CONTENT TO CONVERT:\n"
        f"{raw}"
    )
    repair_task = _Task(
        description=repair_prompt,
        expected_output="One valid JSON object only.",
        agent=agent,
    )
    repair_crew = _Crew(agents=[agent], tasks=[repair_task], verbose=False)
    repair_result = repair_crew.kickoff()
    return _parse_output_to_dict(output_model, repair_result.raw)


def _default_output_for_model(output_model: Any, reason: str) -> dict[str, Any] | None:
    """Return a deterministic schema-valid fallback payload from model defaults."""
    try:
        payload = output_model.model_validate({}).model_dump(mode="json")
        payload["_schema_fallback"] = True
        payload["_schema_fallback_reason"] = reason
        return payload
    except Exception:
        return None

# Pricing per million tokens: {model_prefix: (input_$/M, output_$/M)}
MODEL_PRICING: dict[str, tuple[float, float]] = {
    "anthropic/claude-haiku-4-5": (1.00, 5.00),
    "anthropic/claude-sonnet-4-6": (3.00, 15.00),
    "anthropic/claude-opus-4-6": (15.00, 75.00),
    "anthropic/claude-sonnet-4-20250514": (3.00, 15.00),
    "gpt-4o": (2.50, 10.00),
    "gpt-4o-mini": (0.15, 0.60),
    "gpt-4.1": (2.00, 8.00),
    "gpt-4.1-mini": (0.40, 1.60),
    "gpt-4.1-nano": (0.10, 0.40),
    "o3-mini": (1.10, 4.40),
    "minimax/MiniMax-M2.5": (0.30, 1.20),
}


def _estimate_cost(model: str, prompt_tokens: int, completion_tokens: int) -> float:
    """Estimate USD cost from model name and token counts."""
    pricing = MODEL_PRICING.get(model)
    if not pricing:
        return 0.0
    input_rate, output_rate = pricing
    return (prompt_tokens * input_rate + completion_tokens * output_rate) / 1_000_000


def _write_batch_summary(
    batch_id: str,
    out_dir: Path,
    individual: dict[str, dict[int, Any]],
    ranking_usage: dict[int, dict] | None = None,
    execution_metrics: dict[str, Any] | None = None,
) -> None:
    """Write batch summary with model table and token/cost breakdown."""
    # Aggregate usage across all startups per agent
    aggregated: dict[int, dict] = {}
    fallback_counts: dict[int, int] = {}  # Track number of times each agent used fallback
    
    for _name, outputs in individual.items():
        usage_data = outputs.get("_usage", {})
        for agent_num, info in usage_data.items():
            agent_num = int(agent_num)
            if agent_num not in aggregated:
                aggregated[agent_num] = {
                    "model": info["model"],
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                }
                fallback_counts[agent_num] = 0
            
            aggregated[agent_num]["prompt_tokens"] += info["prompt_tokens"]
            aggregated[agent_num]["completion_tokens"] += info["completion_tokens"]
            aggregated[agent_num]["total_tokens"] += info["total_tokens"]
            
            # Track fallback occurrences
            if info.get("fallback_occurred"):
                fallback_counts[agent_num] += 1

    # Add Agent 7 usage if present
    if ranking_usage:
        for agent_num, info in ranking_usage.items():
            agent_num = int(agent_num)
            aggregated[agent_num] = {
                "model": info["model"],
                "prompt_tokens": info["prompt_tokens"],
                "completion_tokens": info["completion_tokens"],
                "total_tokens": info["total_tokens"],
            }

    lines = [f"# Batch Summary — {batch_id}", ""]
    
    # Retry & Fallback Statistics
    retry_stats = get_retry_stats(batch_id)
    if retry_stats["total_events"] > 0:
        lines.append("## API Retry & Fallback Summary")
        lines.append("")
        lines.append(f"**Total Retry Events**: {retry_stats['total_events']}")
        lines.append(f"**Fallback Events**: {retry_stats['fallback_count']}")
        lines.append(f"**Recovery Events**: {retry_stats['recovery_count']}")
        lines.append(f"**Total Retries**: {retry_stats['total_retries']}")
        lines.append(f"**Average Retries per Event**: {retry_stats['avg_retries']:.1f}")
        lines.append("")
        
        if retry_stats["per_agent"]:
            lines.append("### Per-Agent Breakdown")
            lines.append("")
            lines.append("| Agent | Events | Fallbacks | Recoveries | Total Retries | Intended → Actual Model |")
            lines.append("|-------|--------|-----------|------------|---------------|-------------------------|")
            for agent_stat in retry_stats["per_agent"]:
                agent_num = agent_stat["agent_number"]
                events = agent_stat["events"]
                fallbacks = agent_stat["fallbacks"]
                recoveries = agent_stat["recoveries"]
                total_retries = agent_stat["total_retries"]
                intended = agent_stat["intended_model"]
                actual = agent_stat["actual_model"]
                model_change = f"{intended} → {actual}" if intended != actual else intended
                lines.append(
                    f"| {agent_num}     | {events}      | {fallbacks}         | {recoveries}          | {total_retries}             | {model_change} |"
                )
            lines.append("")
        
        if retry_stats["error_types"]:
            lines.append("### Error Types")
            lines.append("")
            for error_stat in retry_stats["error_types"]:
                lines.append(f"- **{error_stat['error_type']}**: {error_stat['count']} occurrences")
            lines.append("")
    
    lines.append("## Models Used")
    lines.append("")
    lines.append("| Agent | Role                           | Model                              | Notes |")
    lines.append("|-------|--------------------------------|------------------------------------|-------|")
    for agent_num in sorted(aggregated.keys()):
        role = AGENT_ROLES.get(agent_num, "Unknown")
        model = aggregated[agent_num]["model"]
        
        # Add fallback indicator if this agent used fallback
        notes = ""
        if fallback_counts.get(agent_num, 0) > 0:
            notes = f"⚠️ Fallback used {fallback_counts[agent_num]}x"
        
        lines.append(f"| {agent_num}     | {role:<30} | {model:<34} | {notes} |")
    lines.append("")

    # Token usage & cost table
    lines.append("## Token Usage & Cost")
    lines.append("")
    lines.append("| Agent | Prompt Tokens | Completion Tokens | Total Tokens | Cost     |")
    lines.append("|-------|---------------|-------------------|--------------|----------|")
    total_cost = 0.0
    for agent_num in sorted(aggregated.keys()):
        info = aggregated[agent_num]
        cost = _estimate_cost(info["model"], info["prompt_tokens"], info["completion_tokens"])
        total_cost += cost
        lines.append(
            f"| {agent_num}     | {info['prompt_tokens']:,}".ljust(22)
            + f"| {info['completion_tokens']:,}".ljust(22)
            + f"| {info['total_tokens']:,}".ljust(15)
            + f"| ${cost:.3f}".ljust(11) + "|"
        )
    lines.append(
        f"| **Total** |               |                   |              | **${total_cost:.2f}** |"
    )
    lines.append("")

    (out_dir / "batch_summary.md").write_text("\n".join(lines), encoding="utf-8")


def _fmt_list(items: list | str | None, bullet: str = "- ") -> str:
    """Format a list or string as bulleted markdown lines."""
    if not items:
        return f"{bullet}N/A"
    if isinstance(items, str):
        return f"{bullet}{items}"
    return "\n".join(f"{bullet}{item}" for item in items)


def _parse_raw_output(d: dict) -> dict:
    """Parse raw_output JSON string if present, otherwise return dict as-is."""
    raw = d.get("raw_output")
    if raw and isinstance(raw, str):
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return d  # Return original if not valid JSON
    return d


def _write_startup_report(
    startup_name: str, agent_outputs: dict[int | str, Any], path: Path
) -> None:
    """Write a human-readable markdown report for a single startup."""
    a1 = _parse_raw_output(agent_outputs.get(1) or agent_outputs.get("1") or {})
    a2 = agent_outputs.get(2) or agent_outputs.get("2") or {}
    a3 = agent_outputs.get(3) or agent_outputs.get("3") or {}
    a4 = agent_outputs.get(4) or agent_outputs.get("4") or {}
    a5 = _parse_raw_output(agent_outputs.get(5) or agent_outputs.get("5") or {})
    a6 = agent_outputs.get(6) or agent_outputs.get("6") or {}
    tags = agent_outputs.get("_tags") or []

    def g(d: dict, key: str, fallback: str = "N/A") -> str:
        v = d.get(key)
        return str(v) if v is not None else fallback

    lines: list[str] = []

    # Header
    lines.append(f"# {startup_name}")
    lines.append("")
    if a1.get("one_line_description"):
        lines.append(f"> {a1['one_line_description']}")
        lines.append("")

    verdict = g(a2, "verdict")
    weighted = g(a2, "weighted_total_score")
    tier = g(a2, "score_tier")
    recommendation = g(a6, "recommendation")
    lines.append(f"**Verdict:** {verdict} | **Score:** {weighted}/80 ({tier}) | **Recommendation:** {recommendation}")
    lines.append("")

    if tags:
        lines.append("**Tags:** " + ", ".join(f"`{t}`" for t in tags))
        lines.append("")

    lines.append("---")
    lines.append("")

    # --- Agent 1: Startup Brief ---
    lines.append("## 1. Startup Brief (Agent 1)")
    lines.append("")
    for field, label in [
        ("problem", "Problem"),
        ("solution", "Solution"),
        ("target_customer", "Target Customer"),
        ("buyer", "Buyer"),
        ("market", "Market"),
        ("business_model", "Business Model"),
        ("competitors", "Competitors"),
        ("traction", "Traction"),
        ("team", "Team"),
        ("why_now", "Why Now"),
        ("vision", "Vision"),
        ("unfair_advantage", "Unfair Advantage"),
        ("risks", "Risks"),
    ]:
        lines.append(f"- **{label}:** {g(a1, field)}")

    missing = a1.get("missing_info")
    if missing:
        lines.append(f"- **Missing Info:**")
        for item in (missing if isinstance(missing, list) else [missing]):
            lines.append(f"  - {item}")

    inconsistencies = a1.get("inconsistencies")
    if inconsistencies:
        lines.append(f"- **Inconsistencies:**")
        for item in (inconsistencies if isinstance(inconsistencies, list) else [inconsistencies]):
            lines.append(f"  - {item}")

    lines.append(f"- **Clarity Score:** {g(a1, 'clarity_score')}/10")
    lines.append("")

    # --- Agent 2: Venture Analysis ---
    lines.append("## 2. Venture Analysis (Agent 2)")
    lines.append("")
    lines.append(f"**Summary:** {g(a2, 'summary')}")
    lines.append("")

    lines.append("### Scores")
    lines.append("")
    lines.append("| Category | Score |")
    lines.append("|----------|-------|")
    score_fields = [
        ("score_problem_severity", "Problem Severity"),
        ("score_market_size", "Market Size"),
        ("score_differentiation", "Differentiation"),
        ("score_customer_clarity", "Customer Clarity"),
        ("score_founder_insight", "Founder Insight"),
        ("score_business_model", "Business Model"),
        ("score_moat_potential", "Moat Potential"),
        ("score_venture_potential", "Venture Potential"),
        ("score_competition_difficulty", "Competition Difficulty"),
        ("score_execution_feasibility", "Execution Feasibility"),
    ]
    for field, label in score_fields:
        lines.append(f"| {label} | {g(a2, field)}/10 |")
    lines.append(f"| **Total** | **{g(a2, 'total_score')}** |")
    lines.append("")
    lines.append(f"**Weighted Total:** {weighted}/80 — {tier}")
    lines.append("")

    swot = a2.get("swot") or {}
    if swot:
        lines.append("### SWOT")
        lines.append("")
        for category in ["strengths", "weaknesses", "opportunities", "threats"]:
            items = swot.get(category, [])
            lines.append(f"- **{category.title()}:**")
            for item in (items if isinstance(items, list) else [items] if items else []):
                lines.append(f"  - {item}")
        lines.append("")

    reject = a2.get("reject_signals")
    if reject:
        lines.append("**Reject Signals:** " + ", ".join(str(s) for s in reject))
        lines.append("")

    lines.append(f"**Verdict:** {verdict}")
    lines.append("")
    lines.append(f"**Explanation:** {g(a2, 'explanation')}")
    lines.append("")

    # --- Agent 3: Market & Competition ---
    lines.append("## 3. Market & Competition (Agent 3)")
    lines.append("")
    for field, label in [
        ("market_category", "Market Category"),
        ("size_class", "Market Size"),
        ("trend", "Trend"),
        ("direct_competitors", "Direct Competitors"),
        ("indirect_competitors", "Indirect Competitors"),
        ("big_tech_risk", "Big Tech Risk"),
        ("crowdedness", "Crowdedness"),
        ("wedge", "Wedge"),
    ]:
        lines.append(f"- **{label}:** {g(a3, field)}")
    lines.append(f"- **Attractiveness Score:** {g(a3, 'attractiveness_score')}/10")
    lines.append(f"- **Competition Score:** {g(a3, 'competition_score')}/10")
    lines.append(f"- **Conclusion:** {g(a3, 'conclusion')}")
    lines.append("")

    # --- Agent 4: Product & Positioning ---
    lines.append("## 4. Product & Positioning (Agent 4)")
    lines.append("")
    for field, label in [
        ("product_reality", "Product Reality"),
        ("value_prop", "Value Prop"),
        ("killer_feature", "Killer Feature"),
        ("why_care", "Why Care"),
        ("why_not_care", "Why Not Care"),
        ("feature_vs_company", "Feature vs Company"),
        ("wrapper_risk", "Wrapper Risk"),
        ("wedge", "Wedge"),
        ("moat", "Moat"),
        ("positioning", "Positioning"),
        ("six_month_focus", "6-Month Focus"),
    ]:
        lines.append(f"- **{label}:** {g(a4, field)}")
    lines.append("")

    # --- Agent 5: Founder Fit ---
    lines.append("## 5. Founder Fit (Agent 5)")
    lines.append("")
    for field, label in [
        ("founder_fit", "Founder Fit"),
        ("domain", "Domain"),
        ("technical", "Technical"),
        ("distribution", "Distribution"),
        ("strategy", "Strategy"),
        ("ambition", "Ambition"),
        ("execution", "Execution"),
    ]:
        lines.append(f"- **{label}:** {g(a5, field)}")

    missing_roles = a5.get("missing_roles")
    if missing_roles:
        lines.append("- **Missing Roles:**")
        for role in (missing_roles if isinstance(missing_roles, list) else [missing_roles]):
            lines.append(f"  - {role}")

    risks_5 = a5.get("risks")
    if risks_5:
        lines.append("- **Risks:**")
        for risk in (risks_5 if isinstance(risks_5, list) else [risks_5]):
            lines.append(f"  - {risk}")

    lines.append(f"- **Fit Score:** {g(a5, 'fit_score')}/10 | **Execution Score:** {g(a5, 'execution_score')}/10")
    lines.append(f"- **Conclusion:** {g(a5, 'conclusion')}")
    lines.append("")

    # --- Agent 6: Recommendation ---
    lines.append("## 6. Recommendation (Agent 6)")
    lines.append("")
    lines.append(f"**Recommendation:** {recommendation}")
    lines.append("")
    lines.append(f"**Target Customer Segment:** {g(a6, 'customer_segment')}")
    lines.append("")
    lines.append(f"**Wedge:** {g(a6, 'wedge')}")
    lines.append("")

    remove = a6.get("remove")
    if remove:
        lines.append("### Stop Doing")
        lines.append("")
        lines.append(_fmt_list(remove))
        lines.append("")

    emphasize = a6.get("emphasize")
    if emphasize:
        lines.append("### Start Emphasizing")
        lines.append("")
        lines.append(_fmt_list(emphasize))
        lines.append("")

    pivots = a6.get("pivots")
    if pivots:
        lines.append("### Pivot Options")
        lines.append("")
        if isinstance(pivots, list):
            for i, p in enumerate(pivots, 1):
                lines.append(f"{i}. {p}")
        else:
            lines.append(f"1. {pivots}")
        lines.append("")

    if a6.get("positioning_rewrite"):
        lines.append("### Positioning Rewrite")
        lines.append("")
        lines.append(g(a6, "positioning_rewrite"))
        lines.append("")

    if a6.get("thirty_day_plan"):
        lines.append("### 30-Day Plan")
        lines.append("")
        lines.append(g(a6, "thirty_day_plan"))
        lines.append("")

    if a6.get("ninety_day_plan"):
        lines.append("### 90-Day Plan")
        lines.append("")
        lines.append(g(a6, "ninety_day_plan"))
        lines.append("")

    if a6.get("mistake_to_avoid"):
        lines.append("### Mistake to Avoid")
        lines.append("")
        lines.append(g(a6, "mistake_to_avoid"))
        lines.append("")

    path.write_text("\n".join(lines), encoding="utf-8")


def export_results(
    batch_id: str,
    individual: dict[str, dict[int, Any]],
    ranking: dict[str, Any] | None = None,
    ranking_usage: dict[int, dict] | None = None,
    execution_metrics: dict[str, Any] | None = None,
    is_batch_mode: bool = False,
) -> Path:
    """Write pipeline results as JSON files.
    
    Output locations:
    - Batch mode: output/Batch/<batch_id>/
    - Single mode: output/<batch_id>/
    """
    if is_batch_mode:
        out_dir = PROJECT_ROOT / "output" / "Batch" / batch_id
    else:
        out_dir = PROJECT_ROOT / "output" / batch_id
    out_dir.mkdir(parents=True, exist_ok=True)

    for startup_name, agent_outputs in individual.items():
        safe_name = _sanitize_filename(startup_name)
        startup_dir = out_dir / safe_name
        startup_dir.mkdir(parents=True, exist_ok=True)
        (startup_dir / f"{safe_name}.json").write_text(
            json.dumps(agent_outputs, indent=2, default=str),
            encoding="utf-8",
        )
        _write_startup_report(startup_name, agent_outputs, startup_dir / f"{safe_name}.md")

    if ranking is not None:
        (out_dir / "ranking.json").write_text(
            json.dumps(ranking, indent=2, default=str),
            encoding="utf-8",
        )

    _write_batch_summary(batch_id, out_dir, individual, ranking_usage, execution_metrics)

    return out_dir


def _show_agent_timer(
    agent_number: int,
    model_name: str,
    stop_event: threading.Event,
    current_round: int | None = None,
    total_rounds: int | None = None,
    inner_attempt: int | None = None,
    total_inner_attempts: int | None = None,
) -> None:
    """Show a live-updating elapsed timer for a running agent."""
    role = AGENT_ROLES.get(agent_number, "Unknown")
    start = time.time()
    previous_len = 0

    # Assign a unique color per agent for easy identification
    agent_colors = [
        Colors.BRIGHT_CYAN,   # Agent 1
        Colors.BRIGHT_MAGENTA, # Agent 2
        Colors.BRIGHT_YELLOW,  # Agent 3
        Colors.BRIGHT_BLUE,    # Agent 4
        Colors.BRIGHT_RED,     # Agent 5
        Colors.BRIGHT_GREEN,   # Agent 6
        Colors.BRIGHT_CYAN,    # Agent 7
    ]
    agent_color = agent_colors[agent_number % len(agent_colors)]

    def _fit_to_terminal(text: str) -> str:
        # Keep the timer on a single terminal row to prevent wrapped "new lines".
        cols = shutil.get_terminal_size(fallback=(120, 24)).columns
        usable = max(20, cols - 1)
        return text[:usable]

    while not stop_event.is_set():
        elapsed = int(time.time() - start)
        mins, secs = divmod(elapsed, 60)
        round_ctx = ""
        inner_ctx = ""
        if current_round is not None and total_rounds is not None:
            round_ctx = f" | Round {current_round}/{total_rounds}"
        if inner_attempt is not None and total_inner_attempts is not None:
            inner_ctx = f" | Inner {inner_attempt}/{total_inner_attempts}"
        
        # Use color for agent number and role
        msg = f"    {_colorize('⏱', Colors.DIM)} {_colorize(f'Agent {agent_number}', agent_color)} ({_colorize(role, Colors.DIM)}) [{model_name}]{round_ctx}{inner_ctx} ... {mins}m {secs}s"
        msg = _fit_to_terminal(msg)
        # Overwrite in-place and erase any leftover chars from the previous tick.
        clear_tail = " " * max(0, previous_len - len(msg))
        sys.stdout.write(f"\r{msg}{clear_tail}")
        sys.stdout.flush()
        previous_len = len(msg)
        time.sleep(1)

    # Clear the timer line once on exit.
    sys.stdout.write("\r" + " " * previous_len + "\r")
    sys.stdout.flush()


def _run_single_agent(
    agent_number: int,
    submission_text: str,
    prior_context: dict[int, Any] | None = None,
    current_round: int | None = None,
    total_rounds: int | None = None,
    inner_attempt: int | None = None,
    total_inner_attempts: int | None = None,
) -> tuple[dict, dict]:
    """Run a single agent in a one-agent CrewAI Crew and return (output_dict, usage_dict).

    Uses the standard create_task() for Agents 1-6.
    """
    from crewai import Crew, Task

    from src.agents import create_agent
    from src.config import get_model_for_agent
    from src.models import AGENT_OUTPUT_MODELS
    from src.tasks import create_task

    model_name = get_model_for_agent(agent_number)
    role = AGENT_ROLES.get(agent_number, "Unknown")
    agent = create_agent(agent_number)

    task = create_task(
        agent_number=agent_number,
        agent=agent,
        submission_text=submission_text,
        prior_context=prior_context,
    )

    # MiniMax can fail with Instructor "multiple tool calls" on structured output.
    if not _supports_structured_output(model_name):
        task = Task(
            description=task.description,
            expected_output=task.expected_output,
            agent=agent,
        )

    crew = Crew(agents=[agent], tasks=[task], verbose=False)

    # Start live timer
    stop_event = threading.Event()
    agent_start = time.time()
    timer_thread = threading.Thread(
        target=_show_agent_timer,
        args=(
            agent_number,
            model_name,
            stop_event,
            current_round,
            total_rounds,
            inner_attempt,
            total_inner_attempts,
        ),
        daemon=True,
    )
    timer_thread.start()

    def _usage_parts(u: Any) -> tuple[int, int, int]:
        return (
            int(getattr(u, "prompt_tokens", 0) or 0),
            int(getattr(u, "completion_tokens", 0) or 0),
            int(getattr(u, "total_tokens", 0) or 0),
        )

    def _kickoff_with_retry(crew_obj: Any, max_retries: int = 3) -> Any:
        """Execute crew.kickoff() with exponential backoff retry for transient network errors."""
        for attempt in range(max_retries):
            try:
                return crew_obj.kickoff()
            except Exception as e:
                exc_msg_l = str(e).lower()
                is_transient = (
                    "connection reset by peer" in exc_msg_l
                    or "connectionerror" in exc_msg_l
                    or "apiconnectionerror" in type(e).__name__.lower()
                    or "readtimeout" in exc_msg_l
                    or "readerror" in type(e).__name__.lower()
                )
                if not is_transient or attempt == max_retries - 1:
                    raise
                wait_secs = 2 ** attempt
                print(f"\n    ⚠ Transient network error for Agent {agent_number} (attempt {attempt + 1}/{max_retries}), retrying in {wait_secs}s...")
                time.sleep(wait_secs)

    usage_acc = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    result = None
    try:
        result = _kickoff_with_retry(crew)
        p, c, t = _usage_parts(result.token_usage)
        usage_acc["prompt_tokens"] += p
        usage_acc["completion_tokens"] += c
        usage_acc["total_tokens"] += t
    except Exception as exc:
        # Some providers (e.g. MiniMax) trigger "multiple tool calls" errors
        # with structured output. Retry without Pydantic output and parse raw.
        exc_msg = str(exc)
        exc_msg_l = exc_msg.lower()
        if (
            "multiple tool calls" in exc_msg_l
            or "instructor does not support multiple tool calls" in exc_msg_l
            or "instructorretryexception" in type(exc).__name__.lower()
        ):
            print(f"\n    ⚠ Structured output failed for Agent {agent_number}, retrying with raw output...")
            task_fallback = Task(
                description=task.description,
                expected_output=task.expected_output,
                agent=agent,
            )
            crew_fallback = Crew(agents=[agent], tasks=[task_fallback], verbose=False)
            result = _kickoff_with_retry(crew_fallback)
            p, c, t = _usage_parts(result.token_usage)
            usage_acc["prompt_tokens"] += p
            usage_acc["completion_tokens"] += c
            usage_acc["total_tokens"] += t
        else:
            raise
    finally:
        stop_event.set()
        timer_thread.join(timeout=2)

    # Report completion time
    duration = time.time() - agent_start
    mins, secs = divmod(int(duration), 60)
    print(f"    ✓ Agent {agent_number} ({role}) completed in {mins}m {secs}s")

    # Extract output
    output_model = AGENT_OUTPUT_MODELS[agent_number]
    if result.pydantic:
        output_dict = result.pydantic.model_dump(mode="json")
    else:
        parsed_dict = _parse_output_to_dict(output_model, result.raw)
        if parsed_dict is not None:
            output_dict = parsed_dict
        else:
            reason = _parse_failure_reason(output_model, result.raw)
            repaired = None
            try:
                repaired = _repair_output_to_json(agent, output_model, result.raw)
            except Exception:
                repaired = None

            if repaired is not None:
                print(f"    ✓ Agent {agent_number} JSON repaired successfully")
                output_dict = repaired
            else:
                print(f"    ⚠ Could not parse Agent {agent_number} output into expected schema; attempting JSON repair...")
                if reason:
                    print(f"    ↳ Parse diagnostic: {reason}")
                # High-impact agents (1, 2, 6): do one strict rerun before giving up.
                if agent_number in {1, 2, 6}:
                    print(f"    ⚠ JSON repair failed for Agent {agent_number}; running strict format rerun...")
                    field_names = ", ".join(output_model.model_fields.keys())
                    strict_task = Task(
                        description=(
                            task.description
                            + "\n\nCRITICAL OUTPUT REQUIREMENT:\n"
                            + "Return EXACTLY one valid JSON object and nothing else.\n"
                            + "Do not include markdown, backticks, commentary, or prose.\n"
                            + f"Required top-level keys: {field_names}.\n"
                            + "All required fields must be present and valid."
                        ),
                        expected_output=task.expected_output,
                        agent=agent,
                    )
                    strict_crew = Crew(agents=[agent], tasks=[strict_task], verbose=False)
                    strict_result = None
                    try:
                        strict_result = strict_crew.kickoff()
                        p, c, t = _usage_parts(strict_result.token_usage)
                        usage_acc["prompt_tokens"] += p
                        usage_acc["completion_tokens"] += c
                        usage_acc["total_tokens"] += t
                    except Exception:
                        strict_result = None

                    strict_parsed = None
                    if strict_result is not None:
                        if strict_result.pydantic:
                            strict_parsed = strict_result.pydantic.model_dump(mode="json")
                        else:
                            strict_parsed = _parse_output_to_dict(output_model, strict_result.raw)
                            if strict_parsed is None:
                                try:
                                    strict_parsed = _repair_output_to_json(agent, output_model, strict_result.raw)
                                except Exception:
                                    strict_parsed = None

                    if strict_parsed is not None:
                        print(f"    ✓ Agent {agent_number} strict rerun produced valid JSON")
                        output_dict = strict_parsed
                    else:
                        print(f"    ⚠ Agent {agent_number} strict rerun also failed")
                        fallback = _default_output_for_model(
                            output_model,
                            reason=f"Agent {agent_number} parse/repair/rerun failed",
                        )
                        if fallback is not None:
                            print(f"    ✓ Agent {agent_number} using deterministic schema fallback")
                            output_dict = fallback
                        else:
                            output_dict = {"raw_output": result.raw, "_parse_failed": True}
                else:
                    print(f"    ⚠ JSON repair failed for Agent {agent_number}")
                    output_dict = {"raw_output": result.raw, "_parse_failed": True}

    # Token usage
    usage_dict = {
        "model": model_name,
        "prompt_tokens": usage_acc["prompt_tokens"],
        "completion_tokens": usage_acc["completion_tokens"],
        "total_tokens": usage_acc["total_tokens"],
    }

    return output_dict, usage_dict


def main() -> None:
    _ensure_supported_python()
    from src.pipeline import run_batch, run_single

    args = sys.argv[1:]

    if not args:
        # Default: process CourseDocs
        print("No arguments — processing CourseDocs as a single submission.\n")
        submission = load_submission()
        batch_id = _next_batch_id()
        result = run_single(
            startup_name="CourseDocs Startup",
            submission_text=submission,
            batch_id=batch_id,
        )
        print("\n\nFINAL RESULTS")
        print("=" * 60)
        for agent_num in sorted(k for k in result.keys() if isinstance(k, int)):
            print(f"\n--- Agent {agent_num} ---")
            print(json.dumps(result[agent_num], indent=2, default=str))
        if "_tags" in result:
            print(f"\n--- Tags ---")
            print(json.dumps(result["_tags"], indent=2, default=str))

        out_dir = export_results(batch_id, {"CourseDocs Startup": result})
        print(f"\nResults saved to: {out_dir}")
        return

    mode = args[0]

    if mode == "single":
        if len(args) < 2:
            print("Usage: python main.py single <submission_file>")
            sys.exit(1)
        path = Path(args[1])
        if not path.exists():
            print(f"File not found: {path}")
            sys.exit(1)
        submission = load_submission(path)
        name = _extract_startup_name(submission)
        batch_id = _next_batch_id()
        result = run_single(startup_name=name, submission_text=submission, batch_id=batch_id)
        print("\n\nFINAL RESULTS")
        print("=" * 60)
        for agent_num in sorted(result.keys()):
            print(f"\n--- Agent {agent_num} ---")
            print(json.dumps(result[agent_num], indent=2, default=str))

        out_dir = export_results(batch_id, {name: result})
        print(f"\nResults saved to: {out_dir}")

    elif mode == "batch":
        if len(args) < 2:
            print("Usage: python main.py batch <directory>")
            sys.exit(1)
        folder = Path(args[1])
        if not folder.is_dir():
            print(f"Not a directory: {folder}")
            sys.exit(1)

        submissions: dict[str, str] = {}
        subdirs = sorted([d for d in folder.iterdir() if d.is_dir() and not d.name.startswith(".")])
        
        def read_file_content(f: Path) -> str:
            """Read text content from various file types."""
            ext = f.suffix.lower()
            if ext == '.md' or ext == '.txt':
                return f.read_text(encoding="utf-8", errors="replace")
            elif ext == '.docx':
                from docx import Document
                doc = Document(f)
                return "\n".join([p.text for p in doc.paragraphs])
            elif ext == '.pdf':
                from PyPDF2 import PdfReader
                try:
                    reader = PdfReader(f)
                    return "\n".join([page.extract_text() or "" for page in reader.pages])
                except Exception:
                    return ""
            else:
                return f.read_text(encoding="utf-8", errors="replace")
        
        if subdirs:
            # New structure: each subfolder is one startup
            for subdir in subdirs:
                # Accept .md, .txt, .docx, .pdf files
                allowed_exts = {'.md', '.txt', '.docx', '.pdf'}
                all_files = sorted([f for f in subdir.iterdir() 
                                   if f.is_file() and not f.name.startswith(".")
                                   and f.suffix.lower() in allowed_exts])
                if not all_files:
                    continue
                combined = "\n\n".join(read_file_content(f) for f in all_files)
                submissions[subdir.name] = combined
        else:
            # Legacy: .md files directly in the folder
            for f in sorted(folder.glob("*.md")):
                text = f.read_text(encoding="utf-8")
                name = _extract_startup_name(text)
                submissions[name] = text

        if not submissions:
            print(f"No startup submissions found in {folder}")
            sys.exit(1)

        print(f"Found {len(submissions)} submissions: {list(submissions.keys())}\n")
        batch_id = _next_batch_id()
        result = run_batch(submissions, batch_id=batch_id)

        print("\n\nINDIVIDUAL RESULTS")
        print("=" * 60)
        for name, outputs in result["individual"].items():
            print(f"\n{'#'*40}")
            print(f"  {name}")
            print(f"{'#'*40}")
            for agent_num in sorted(k for k in outputs.keys() if isinstance(k, int)):
                print(f"\n--- Agent {agent_num} ---")
                print(json.dumps(outputs[agent_num], indent=2, default=str))
            if "_tags" in outputs:
                print(f"\n--- Tags ---")
                print(json.dumps(outputs["_tags"], indent=2, default=str))

        if result["ranking"]:
            print("\n\nCOHORT RANKING")
            print("=" * 60)
            print(json.dumps(result["ranking"], indent=2, default=str))

        out_dir = export_results(
            batch_id, result["individual"], result["ranking"], result.get("ranking_usage"),
            is_batch_mode=True,
        )
        print(f"\nResults saved to: {out_dir}")

    else:
        print(f"Unknown mode: {mode}")
        print("Usage:")
        print("  python main.py                     # Process CourseDocs")
        print("  python main.py single <file>       # Process one submission")
        print("  python main.py batch <directory>  # Process multiple + rank")
        sys.exit(1)


if __name__ == "__main__":
    main()
