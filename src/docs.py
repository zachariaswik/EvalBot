"""Load startup submission documents from CourseDocs or an arbitrary path."""

from __future__ import annotations

from pathlib import Path

from .config import COURSE_DOCS_DIR


def _extract_text_from_docx(docx_path: Path) -> str | None:
    """Extract text and tables from DOCX file, preserving document order.
    
    Returns None if extraction fails. See main.py for full implementation details.
    """
    try:
        import docx
        doc = docx.Document(docx_path)
        parts = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                        break
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells)
                            if row_text.strip():
                                parts.append(row_text)
                        parts.append("")
                        break
        
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def load_submission(source: Path | None = None) -> str:
    """Read submission files and return text suitable as input for Agent 1.
    
    Supports:
    - PDF files: Passed as file references for Claude to read directly
    - DOCX files: Text and tables extracted
    - Text files: Read as UTF-8
    - Directories: All supported files concatenated
    """
    folder = source or COURSE_DOCS_DIR
    if folder.is_file():
        # Single file case
        if folder.suffix.lower() == ".pdf":
            # Return PDF file reference for Claude to read directly
            return f"[PDF_FILE: {folder.absolute()}]"
        elif folder.suffix.lower() == ".docx":
            # Extract text and tables from DOCX
            content = _extract_text_from_docx(folder)
            if content:
                return content
            raise ValueError(f"Could not extract content from DOCX: {folder}")
        # Plain text file
        return folder.read_text(encoding="utf-8")

    # Directory case: look for .md, .pdf, or .docx files
    md_files = sorted(folder.glob("*.md"))
    pdf_files = sorted(folder.glob("*.pdf"))
    docx_files = sorted(folder.glob("*.docx"))
    
    if not md_files and not pdf_files and not docx_files:
        raise FileNotFoundError(f"No .md, .pdf, or .docx files found in {folder}")

    parts: list[str] = []
    
    # Handle markdown files
    for f in md_files:
        parts.append(f"# {f.stem}\n\n{f.read_text(encoding='utf-8')}")
    
    # Handle PDF files - pass as file references
    for f in pdf_files:
        parts.append(f"[PDF_FILE: {f.absolute()}]")
    
    # Handle DOCX files - extract text and tables
    for f in docx_files:
        content = _extract_text_from_docx(f)
        if content:
            parts.append(f"# {f.stem}\n\n{content}")
    
    return "\n\n---\n\n".join(parts)
